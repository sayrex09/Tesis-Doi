import csv
import re
import unicodedata
from io import BytesIO, StringIO

import pandas as pd


CAMPO_ALIASES = {
    "title": [
        "title",
        "document title",
        "article title",
        "nombre del articulo",
        "nombre del artículo",
    ],
    "authors": ["authors", "author(s)", "author names", "autor(es)", "autores"],
    "year": ["year", "publication year", "pubyear", "año publicacion", "año publicación"],
    "source_title": [
        "source_title",
        "source title",
        "journal",
        "publication name",
        "revista/conferencia",
        "revista",
        "conferencia",
        "nombre de la revista",
        "nombre de la conferencia",
    ],
    "volume": ["volume", "vol", "volumen"],
    "issue": ["issue", "issues", "number", "numero", "número"],
    "pages": ["pages", "page", "pagination", "page range", "paginas", "páginas"],
    "doi": ["doi", "digital object identifier"],
    "url": ["url", "link", "enlace"],
    "publisher": ["publisher", "editorial"],
    "document_type": ["document type", "tipo documento", "source document type"],
}


def _normalizar_clave(texto):
    texto = str(texto or "").lower().strip()
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(caracter for caracter in texto if not unicodedata.combining(caracter))
    texto = re.sub(r"[^a-z0-9]+", " ", texto)
    return re.sub(r"\s+", " ", texto).strip()


def _esta_vacio(valor):
    if valor is None:
        return True
    try:
        if pd.isna(valor):
            return True
    except (TypeError, ValueError):
        pass
    if isinstance(valor, str) and valor.strip().lower() in {"", "nan", "none", "null", "n/a", "na"}:
        return True
    return False


def _limpiar_texto(valor):
    if _esta_vacio(valor):
        return ""
    texto = str(valor).replace("\ufeff", " ")
    texto = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", texto)
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip()


def _como_diccionario(articulo):
    if hasattr(articulo, "to_dict"):
        return articulo.to_dict()
    if isinstance(articulo, dict):
        return articulo
    return {}


def _obtener_valor(articulo, campo, limpiar=True):
    datos = _como_diccionario(articulo)
    datos_normalizados = {_normalizar_clave(clave): valor for clave, valor in datos.items()}

    for alias in CAMPO_ALIASES[campo]:
        clave = _normalizar_clave(alias)
        if clave in datos_normalizados:
            valor = datos_normalizados[clave]
            return _limpiar_texto(valor) if limpiar else valor
    return ""


def _dividir_autores(autores):
    if _esta_vacio(autores):
        return []

    if isinstance(autores, (list, tuple)):
        partes = []
        for autor in autores:
            if isinstance(autor, dict):
                nombre = _limpiar_texto(autor.get("given", ""))
                apellido = _limpiar_texto(autor.get("family", ""))
                partes.append(f"{nombre} {apellido}".strip())
            else:
                partes.append(_limpiar_texto(autor))
        return [parte for parte in partes if parte]

    texto = _limpiar_texto(autores)
    if not texto:
        return []

    if ";" in texto:
        return [parte.strip() for parte in texto.split(";") if parte.strip()]
    if "|" in texto:
        return [parte.strip() for parte in texto.split("|") if parte.strip()]
    if re.search(r"\s+and\s+", texto, flags=re.IGNORECASE):
        return [parte.strip() for parte in re.split(r"\s+and\s+", texto, flags=re.IGNORECASE) if parte.strip()]

    autores_apellido_iniciales = re.findall(
        r"([^,;]+),\s*((?:[A-ZÁÉÍÓÚÑ]\.?\s*){1,5})(?:,|$)",
        texto,
    )
    if len(autores_apellido_iniciales) >= 2:
        return [f"{apellido.strip()}, {iniciales.strip()}" for apellido, iniciales in autores_apellido_iniciales]
    if "," in texto:
        partes = [parte.strip() for parte in texto.split(",") if parte.strip()]
        if len(partes) > 2 or not re.fullmatch(r"(?:[A-ZÁÉÍÓÚÑ]\.?\s*){1,5}", partes[-1]):
            return partes

    return [texto]


def _iniciales_desde_nombres(nombres):
    iniciales = []
    for token in re.split(r"\s+", _limpiar_texto(nombres)):
        letras = re.findall(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]", token)
        if not letras:
            continue
        if len(letras) > 1 and "." in token:
            iniciales.extend(letra.upper() for letra in letras)
        else:
            iniciales.append(letras[0].upper())
    return " ".join(f"{inicial}." for inicial in iniciales)


def _formatear_autor_ieee(autor):
    autor = _limpiar_texto(autor)
    if not autor:
        return ""

    autor = re.sub(r"\s*\([^)]*\)", "", autor).strip()
    if "," in autor:
        apellido, nombres = autor.split(",", 1)
        iniciales = _iniciales_desde_nombres(nombres)
        apellido = _limpiar_texto(apellido)
        return f"{iniciales} {apellido}".strip() if iniciales else apellido

    partes = autor.split()
    if len(partes) == 1:
        return autor

    apellido = partes[-1]
    nombres = " ".join(partes[:-1])
    iniciales = _iniciales_desde_nombres(nombres)
    return f"{iniciales} {apellido}".strip() if iniciales else autor


def formatear_autores_ieee(autores, usar_et_al=False, limite_et_al=6):
    """Formatea autores como iniciales y apellido siguiendo el estilo IEEE."""
    autores_formateados = [
        autor for autor in (_formatear_autor_ieee(autor) for autor in _dividir_autores(autores)) if autor
    ]

    if not autores_formateados:
        return ""

    if usar_et_al and len(autores_formateados) > int(limite_et_al or 0):
        return f"{autores_formateados[0]} et al."

    if len(autores_formateados) == 1:
        return autores_formateados[0]
    if len(autores_formateados) == 2:
        return f"{autores_formateados[0]} and {autores_formateados[1]}"

    return f"{', '.join(autores_formateados[:-1])}, and {autores_formateados[-1]}"


def _formatear_paginas(paginas):
    paginas = _limpiar_texto(paginas)
    if not paginas:
        return ""

    paginas = paginas.replace("–", "-").replace("—", "-")
    paginas = re.sub(r"^(pp?\.?\s*)", "", paginas, flags=re.IGNORECASE).strip()
    etiqueta = "pp." if re.search(r"\d\s*-\s*\d", paginas) else "p."
    return f"{etiqueta} {paginas}"


def _limpiar_doi_referencia(doi):
    doi = _limpiar_texto(doi)
    if not doi:
        return ""
    doi = re.sub(r"^(https?://(dx\.)?doi\.org/|doi\.org/|dx\.doi\.org/|doi:)", "", doi, flags=re.IGNORECASE)
    return doi.strip().strip(".,;")


def _limpiar_titulo_referencia(titulo):
    return _limpiar_texto(titulo).rstrip(" .,;:")


def generar_cita_ieee(articulo, numero=None, usar_et_al=False, limite_et_al=6):
    """Genera una referencia IEEE individual omitiendo campos bibliográficos vacíos."""
    autores = formatear_autores_ieee(
        _obtener_valor(articulo, "authors", limpiar=False),
        usar_et_al=usar_et_al,
        limite_et_al=limite_et_al,
    )
    titulo = _limpiar_titulo_referencia(_obtener_valor(articulo, "title"))
    fuente = _obtener_valor(articulo, "source_title")
    volumen = _obtener_valor(articulo, "volume")
    numero_revista = _obtener_valor(articulo, "issue")
    paginas = _formatear_paginas(_obtener_valor(articulo, "pages"))
    anio = _obtener_valor(articulo, "year")
    doi = _limpiar_doi_referencia(_obtener_valor(articulo, "doi"))
    url = _obtener_valor(articulo, "url")

    detalles = []
    if fuente:
        detalles.append(fuente)
    if volumen:
        detalles.append(f"vol. {volumen}")
    if numero_revista:
        detalles.append(f"no. {numero_revista}")
    if paginas:
        detalles.append(paginas)
    if anio and anio not in {"0", "0.0"}:
        detalles.append(anio)
    if doi:
        detalles.append(f"doi: {doi}")
    elif url:
        detalles.append(f"Available: {url}")

    partes_iniciales = []
    if autores:
        partes_iniciales.append(autores)
    if titulo:
        titulo_formateado = f"\"{titulo},\"" if detalles else f"\"{titulo}\""
        partes_iniciales.append(titulo_formateado)

    if partes_iniciales:
        cita = ", ".join(partes_iniciales)
    elif detalles:
        cita = detalles.pop(0)
    else:
        cita = "Referencia sin datos bibliográficos suficientes"

    if detalles:
        cita = f"{cita} {', '.join(detalles)}"

    cita = cita.rstrip(" .,") + "."
    return f"[{numero}] {cita}" if numero is not None else cita


def generar_lista_referencias_ieee(articulos, usar_et_al=False, limite_et_al=6):
    """Genera una lista numerada de referencias IEEE desde una lista o DataFrame."""
    if isinstance(articulos, pd.DataFrame):
        registros = articulos.fillna("").to_dict("records")
    else:
        registros = list(articulos or [])

    return [
        generar_cita_ieee(
            articulo,
            numero=indice,
            usar_et_al=usar_et_al,
            limite_et_al=limite_et_al,
        )
        for indice, articulo in enumerate(registros, start=1)
    ]


def exportar_referencias_txt(referencias):
    """Devuelve las referencias IEEE en bytes para descargar como TXT."""
    return ("\n".join(referencias) + "\n").encode("utf-8")


def exportar_referencias_docx(referencias):
    """Devuelve un documento DOCX en memoria con las referencias IEEE."""
    from docx import Document

    documento = Document()
    documento.add_heading("Referencias IEEE", level=1)
    for referencia in referencias:
        documento.add_paragraph(referencia)

    buffer = BytesIO()
    documento.save(buffer)
    buffer.seek(0)
    return buffer


def exportar_referencias_csv(referencias):
    """Devuelve un CSV con número y referencia IEEE."""
    buffer = StringIO()
    writer = csv.writer(buffer)
    writer.writerow(["numero", "referencia_ieee"])
    for indice, referencia in enumerate(referencias, start=1):
        writer.writerow([indice, referencia])
    return buffer.getvalue().encode("utf-8-sig")


def _autores_bibtex(autores):
    partes = [_formatear_autor_bibtex(autor) for autor in _dividir_autores(autores)]
    return " and ".join(parte for parte in partes if parte)


def _formatear_autor_bibtex(autor):
    autor = _limpiar_texto(autor)
    if not autor:
        return ""

    if "," in autor:
        apellido, nombres = autor.split(",", 1)
        apellido = _limpiar_texto(apellido)
        iniciales = _iniciales_desde_nombres(nombres)
        nombres = iniciales or _limpiar_texto(nombres)
        return f"{apellido}, {nombres}".strip(", ")

    partes = autor.split()
    if len(partes) == 1:
        return autor

    apellido = partes[-1]
    nombres = " ".join(partes[:-1])
    iniciales = _iniciales_desde_nombres(nombres)
    return f"{apellido}, {iniciales}".strip(", ") if iniciales else autor


def _clave_bibtex(articulo, indice):
    autores = _dividir_autores(_obtener_valor(articulo, "authors", limpiar=False))
    primer_autor = autores[0] if autores else "ref"
    if "," in primer_autor:
        apellido = primer_autor.split(",", 1)[0]
    else:
        apellido = primer_autor.split()[-1] if primer_autor.split() else "ref"
    apellido = _normalizar_bibtex_key(_apellido_para_clave(apellido)) or "ref"
    anio = re.search(r"\b(19|20)\d{2}\b", _obtener_valor(articulo, "year") or "")
    anio = anio.group(0) if anio else "sfa"
    palabra_titulo = _palabra_clave_titulo(_obtener_valor(articulo, "title"))
    clave = f"{apellido}{anio}{palabra_titulo}"
    return clave or f"ref{indice}"


def _apellido_para_clave(apellido):
    partes = _limpiar_texto(apellido).split()
    if not partes:
        return "ref"

    particulas = {"da", "de", "del", "di", "dos", "la", "las", "le", "los", "van", "von"}
    for parte in partes:
        if _normalizar_clave(parte) not in particulas:
            return parte
    return partes[-1]


def _normalizar_bibtex_key(texto):
    texto = unicodedata.normalize("NFKD", _limpiar_texto(texto))
    texto = "".join(caracter for caracter in texto if not unicodedata.combining(caracter))
    texto = re.sub(r"[^A-Za-z0-9]+", "", texto)
    if not texto:
        return ""
    return texto[0].lower() + texto[1:]


def _palabra_clave_titulo(titulo):
    texto = unicodedata.normalize("NFKD", _limpiar_texto(titulo).lower())
    texto = "".join(caracter for caracter in texto if not unicodedata.combining(caracter))
    palabras_vacias = {
        "a",
        "al",
        "and",
        "de",
        "del",
        "el",
        "en",
        "for",
        "in",
        "la",
        "las",
        "los",
        "of",
        "on",
        "para",
        "the",
        "to",
        "un",
        "una",
        "y",
    }
    for palabra in re.findall(r"[a-z0-9]+", texto):
        if len(palabra) >= 4 and palabra not in palabras_vacias:
            return palabra
    return "referencia"


def _claves_unicas_bibtex(registros):
    conteo = {}
    claves = []
    for indice, articulo in enumerate(registros, start=1):
        clave_base = _clave_bibtex(articulo, indice)
        conteo[clave_base] = conteo.get(clave_base, 0) + 1
        claves.append(clave_base if conteo[clave_base] == 1 else f"{clave_base}{conteo[clave_base]}")
    return claves


def _escapar_bibtex(valor):
    valor = _limpiar_texto(valor)
    return valor.replace("\\", "\\\\")


def _proteger_mayusculas_titulo(titulo):
    titulo = _limpiar_texto(titulo)
    tokens = []
    for token in titulo.split(" "):
        limpio = token.strip(".,:;()[]")
        if (
            len(limpio) >= 2
            and limpio.isupper()
            and any(caracter.isalpha() for caracter in limpio)
            and not token.startswith("{")
        ):
            token = token.replace(limpio, f"{{{limpio}}}", 1)
        tokens.append(token)
    return " ".join(tokens)


def _paginas_bibtex(paginas):
    paginas = _limpiar_texto(paginas)
    if not paginas:
        return ""
    paginas = paginas.replace("–", "-").replace("—", "-")
    paginas = re.sub(r"^(pp?\.?\s*)", "", paginas, flags=re.IGNORECASE).strip()
    return re.sub(r"\s*-\s*", "--", paginas)


def _agregar_campo_bibtex(lineas, campo, valor):
    valor = _limpiar_texto(valor)
    if valor and valor not in {"0", "0.0"}:
        lineas.append(f"  {campo} = {{{_escapar_bibtex(valor)}}},")


def exportar_referencias_bibtex(articulos):
    """Genera entradas BibTeX validas en formato compatible con IEEEtran."""
    if isinstance(articulos, pd.DataFrame):
        registros = articulos.fillna("").to_dict("records")
    else:
        registros = list(articulos or [])

    entradas = []
    claves = _claves_unicas_bibtex(registros)
    for articulo, clave in zip(registros, claves):
        tipo_documento = _obtener_valor(articulo, "document_type").lower()
        es_conferencia = "conferencia" in tipo_documento or "conference" in tipo_documento
        entry_type = "inproceedings" if es_conferencia else "article"

        fuente = _obtener_valor(articulo, "source_title")
        titulo = _proteger_mayusculas_titulo(_obtener_valor(articulo, "title"))
        paginas = _paginas_bibtex(_obtener_valor(articulo, "pages"))
        doi = _limpiar_doi_referencia(_obtener_valor(articulo, "doi"))
        link = _obtener_valor(articulo, "url")

        lineas = [f"@{entry_type}{{{clave},"]

        _agregar_campo_bibtex(lineas, "author", _autores_bibtex(_obtener_valor(articulo, "authors", limpiar=False)))
        _agregar_campo_bibtex(lineas, "title", titulo)
        _agregar_campo_bibtex(lineas, "booktitle" if es_conferencia else "journal", fuente)

        if not es_conferencia:
            _agregar_campo_bibtex(lineas, "volume", _obtener_valor(articulo, "volume"))
            _agregar_campo_bibtex(lineas, "number", _obtener_valor(articulo, "issue"))

        _agregar_campo_bibtex(lineas, "pages", paginas)
        _agregar_campo_bibtex(lineas, "year", _obtener_valor(articulo, "year"))
        _agregar_campo_bibtex(lineas, "doi", doi)
        _agregar_campo_bibtex(lineas, "link", link)

        if len(lineas) > 1:
            lineas[-1] = lineas[-1].rstrip(",")
        lineas.append("}")
        entradas.append("\n".join(lineas))

    return ("\n\n".join(entradas) + "\n").encode("utf-8")


def preparar_exportacion_referencias(articulos, formato, usar_et_al=False, limite_et_al=6):
    """Prepara nombre, MIME y contenido descargable para Streamlit."""
    formato = str(formato or "TXT").upper()
    referencias = generar_lista_referencias_ieee(
        articulos,
        usar_et_al=usar_et_al,
        limite_et_al=limite_et_al,
    )

    if formato == "DOCX":
        return "referencias_ieee.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", exportar_referencias_docx(referencias)
    if formato == "CSV":
        return "referencias_ieee.csv", "text/csv", exportar_referencias_csv(referencias)
    if formato == "BIBTEX":
        return "referencias_ieee.bib", "application/x-bibtex", exportar_referencias_bibtex(articulos)

    return "referencias_ieee.txt", "text/plain", exportar_referencias_txt(referencias)
